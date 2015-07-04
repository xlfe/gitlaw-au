
from lxml import html
import requests
from docx import Document
import os

link_xtree = '//*[@id="ctl00_MainContent_AttachmentsRepeater_ctl00_ArtifactVersionRenderer_Repeater1_ctl00_ArtifactFormatTableRenderer1_RadGridNonHtml_ctl00_ctl04_hlPrimaryDoc"]'

SPIDER_DELAY=10000
CACHE_DIR = '.cache-gitlaw-au'

class Legislation(object):

    def __init__(self, comlaw_id):
        self.cli = comlaw_id

        if not os.path.exists(self.filename):
            print 'Not cached, downloading'
            self.download_legislation_document()
        else:
            print 'Loading from cache'
        self.docx = Document(self.filename)
        self.paragraphs = [p for p in self.docx.paragraphs]
        print 'Loaded {} "{}"'.format(self.cli, self.titles['short'])

    @property
    def filename(self):
        cache_dir = os.path.join(os.path.expanduser('~'), CACHE_DIR)
        if not os.path.exists(cache_dir):
            os.mkdir(cache_dir)
        return os.path.join(cache_dir, '{}.docx'.format(self.cli))

    def download_legislation_document(self):

        download_pg = 'http://www.comlaw.gov.au/Details/{}/Download'.format(self.cli)
        r = requests.get(download_pg, verify=False)

        if r.status_code != 200:
            raise Exception('Error loading details page for {}'.format(self.cli))

        h = html.fromstring(r.content)
        a = h.xpath(link_xtree)[0]
        d = requests.get(a.attrib['href'], verify = False)

        if d.status_code != 200:
            raise Exception('Error downloading file')

        with open(self.filename, 'w') as out:
            out.write(d.content)

        print 'Downloaded {}'.format(self.cli)

    def text(self):
        for p in self.docx.paragraphs:
            if len(p.text.strip()) == 0:
                continue
            print p.style.name
            print p.style
            for i in range(len(p.text)/80 + 1):
                line = p.text[i*80:(i+1)*80]
                print '\t{}'.format(line.encode('ascii','ignore'))

    def get_styles(self, style_nm):
        return filter(lambda k: k.style.name.lower() == style_nm.lower(),self.paragraphs)

    @property
    def titles(self):
        short = self.get_styles('ShortT')
        long = self.get_styles('LongT')
        assert len(short) == 1
        assert len(long) == 1
        return {
            'short': short[0].text,
            'long': long[0].text
        }


def test():
    import collections
    c = collections.Counter()
    l = Legislation('C2015A00061')
    d = l.docx
    for p in d.paragraphs:
        c[p.style.name] += 1

    for name, count in sorted(c.iteritems(),key=lambda k:k[0]):
        print '{:<20} {:,}'.format(name, count)