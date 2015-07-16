
# Convert the downloaded documents to markdown and put them in the right folder

INPUT = './details_current.json'


import collections
import subprocess
import os
import json

from docx import Document
from docx.shared import Pt

import re

code_map = [
    ('\r', ''),     #Character return
    (u'\xa0', ' '),     #Non breaking space
    (u'\u00b0', ' (?degree?) '),     #degree sign
    (u'\u00b4', "'"),     #acute accent
    (u'\u2018', "'"),   #Left single quotation mark
    (u'\u2019', "'"),   #Right single quotation mark
    (u'\u2022', ' '),   #Bullet
    (u'\u201c', '"'),   #Left double quotation mark
    (u'\u201d', '"'),   #Right double quotation mark
    (u'\u2013', '-'),   #En dash
    (u'\u2014', '--'),  #Em dash
    (u'\u2026', '...'), #horizontal elipsis
    (u'\u2011', '-')    #Non-breaking hyphen
]
STYLE = collections.namedtuple('STYLE', 'BOLD ITALIC')

class NEWPARA(object):
    def __repr__(self):
        return '\n'

class TEXT(object):
    ARGS = ['bold', 'italic', 'indent', 'heading', 'text']

    def __init__(self, *args, **kwargs):

        if len(args) == len(self.ARGS):
            for i, arg in enumerate(self.ARGS):
                setattr(self, arg, args[i])
        else:
            for arg in self.ARGS:
                try:
                    setattr(self, arg, kwargs[arg])
                except KeyError:
                    setattr(self, arg, None)

        assert self.text is not None

    def _asdict(self):
        return {k:getattr(self, k) for k in self.ARGS}

    def __repr__(self):
        return str(tuple(getattr(self,k) for k in self.ARGS))

    def join(self, other):
        """try to add two text objects together"""


        #Only check styles if there is non-whitespace characters
        if other.text.strip() != '':
            for arg in self.ARGS[:-1]:
                if getattr(self, arg) != getattr(other, arg):
                    raise ValueError('{} is not equal'.format(arg))

        added = TEXT(**self._asdict())
        added.text += other.text
        return added

def para_indent(p):
    indent = p.paragraph_format.left_indent.pt if p.paragraph_format.left_indent is not None else 0
    if not indent:
        indent = p.style.paragraph_format.left_indent.pt if p.style.paragraph_format.left_indent is not None else 0
    return indent

def pts_to_header(pt):
    if pt > 18:
        return 1

    elif pt > 16:
        return 2

    elif pt > 14:
        return 3

    elif pt > 12:
        return 4

    return 0


class TextUtil(object):

    PANDOC_OPTIONS = 'pandoc -f docx -t markdown_github --no-wrap --ascii '
    def check_dir(self, dir):

        if not os.path.exists(dir):
            os.makedirs(dir)

    def convert(self, input, output, type):


        if type == 'iconDOCX':
            print 'Opening {}'.format(input)
            docx = Document(input)
        else:
            print 'converting from {} {}'.format(type, input)
            command = 'textutil -convert docx -stdout ./{} >| ./tmp'.format(input)
            subprocess.check_call(command, shell=True)
            docx = Document('./tmp')
            # os.remove('./tmp')

        normal = docx.styles['Normal']

        if docx.styles['Normal'] and docx.styles['Normal'].font and docx.styles['Normal'].font.size:
            normal_size = docx.styles['Normal'].font.size.pt
        else:
            normal_size = Pt(11).pt


        indents = collections.defaultdict(int)
        font_sizes = collections.defaultdict(int)
        prev_indent = 0

        for p in docx.paragraphs:
            if not p.text.strip():
                continue

            indent = para_indent(p)
            indents[indent] += 1

            if p.style.font and p.style.font.size:
                p_size = p.style.font.size.pt
            elif p.style.base_style and p.style.base_style.font and p.style.base_style.font.size:
                p_size = p.style.base_style.font.size.pt
            else:
                p_size = normal_size

            for r in p.runs:

                r_size = r.font.size.pt if r.font.size is not None else p_size
                font_sizes[r_size] += 1

        font_sizes_ordered = sorted(font_sizes.iterkeys(), key=lambda k:int(k), reverse=True)
        font_sizes_usage = sorted(font_sizes.iteritems(), key=lambda k:k[1], reverse=True)
        most_common_size = font_sizes_usage[0][0]
        idx_most_common = font_sizes_ordered.index(most_common_size)
        out = []

        for p in docx.paragraphs:

            if not p.text.strip():
                continue


            if p.style.font and p.style.font.size:
                p_size = p.style.font.size.pt
            elif p.style.base_style and p.style.base_style.font and p.style.base_style.font.size:
                p_size = p.style.base_style.font.size.pt
            else:
                p_size = normal_size

            _indent = para_indent(p)

            if _indent > 0:
                indent = int(_indent / p_size)
            else:
                indent = 0

            sizes = collections.defaultdict(int)
            for r in p.runs:

                r_size = r.font.size.pt if r.font.size is not None else p_size

                size_index = font_sizes_ordered.index(r_size) + 1

                if size_index > 6 or size_index >= idx_most_common:
                    size_index = 0
                sizes[size_index] += 1

            size_index = sorted(sizes.iteritems(), key=lambda k:k[1], reverse=True)[0][0]

            for r in p.runs:

                run_indent = 0

                text = r.text
                for frm, to in code_map:
                    text = text.replace(frm, to)

                #Remove EMBED WORD PICTURE
                if text.strip().startswith('EMBED'):
                    continue

                remove = r'(DOCPROPERTY ([a-zA-Z]+)|TOC \\o|PAGEREF _[a-zA-Z0-9]+ \\h [\d]+|STYLEREF [a-zA-Z0-9]+)'
                text = re.sub(remove, '', text)

                out.append(TEXT(
                    1 if r.bold == True and size_index == 0 else 0,
                    1 if r.italic == True and size_index == 0 else 0,
                    indent + run_indent,
                    size_index,
                    text
                ))

            #Paragraph
            out.append(NEWPARA())


        return out


class StyleManager(object):

    def __init__(self):
        self.s = STYLE(False, False)

    def this_style(self, bold, italic, text):

        out = ''

        ts = STYLE(bold, italic)

        if self.s != ts:
            #change in style

            out += self.close()

            if ts.BOLD:
                out += '**'

            if ts.ITALIC:
                out += '_'

        self.s = ts
        out += text

        return out


    def close(self):
        out = ''
        if self.s.ITALIC:
            out += '_'
        if self.s.BOLD:
            out += '**'
        return out


def join_styles(input):
    """merge text from styles which are next to each other and identical"""

    joined = []
    prev = None

    for i in input:

        if isinstance(i, NEWPARA):

            if prev:
                joined.append(prev)

            joined.append(i)
            prev = None
            continue

        if prev:
            try:
                i = prev.join(i)
            except ValueError:
                joined.append(prev)

        prev = i

    return joined


no_style = STYLE(False, False)
def apply_bold_italic(joined):
    """
        Apply bold and italic rules to text
    """

    style = StyleManager()
    last = None

    for j in joined:

        if isinstance(j, NEWPARA):
            if isinstance(last, TEXT):
                last.text += style.close()
            last = j
            continue

        j.text = style.this_style(j.bold, j.italic, j.text)
        j.bold = False
        j.italic = False
        last = j


    if isinstance(last, TEXT):
        last.text += style.close()
    return joined


def convert_indentation(joined):

    for j in joined:
        if isinstance(j, NEWPARA):
            continue

        #Move spaces to before (i) etc
        j.text = re.sub(r'^([\s]*)([a-z0-9\sA-Z]+:|\([0-9a-zA-Z\.]+\))([\s]*)', r'\1\3\2 ', j.text, count=1)

        while j.text.startswith('\t'):
            # if j.text.startswith(' '):
            #     j.indent += 1
            # elif j.text.startswith('\t'):
            j.indent += 4
            j.text = j.text[1:]

    return joined


def apply_indentation(joined):

    output = []

    for j in joined:

        if isinstance(j,NEWPARA):
            output.append('\n')
            output.append('\n')
            continue

        has_nl = '\n' in j.text

        for text in j.text.split('\n'):

            #Contents often has    1    Name   2
            text = re.sub(r'^[\s]*([\d]+)\t([^\t]+)\t([\d]+)', r'\1 \2 ', text)


            #only apply headings or indents after a newline...

            if j.heading > 0:
                output.append('#' * j.heading)

                # Remove spaces from the front of a heading...
                text = re.sub(r'^[\s]*',' ',text, count=1)

            elif j.indent > 0:
                output.append(' ' * j.indent + ' * ')

            output.append(text)
            if has_nl:
                output.append('\n')




    return ''.join(output)






with open(INPUT, 'r') as inp:

    details = json.load(inp)



# {
#     u'status': u'Current',
#     u'uuid': u'2f8dc4ff-7d4d-4769-9886-e9d93845ba03',
#     u'title': u'Excise Act 1901',
#     u'ComLawID': u'C2015C00135',
#     u'subname': u'Act Compilation',
#     u'volname': u'',
#     u'type': u'iconDOCX',
#     u'pages': u''
# }



converter = TextUtil()


import re
keepcharacters = (' ','.','-','(',')','/')
WRITE = True

print 'Loaded {} documents'.format(len(details))

for doc in details:

    try:
        if int(doc['pages'].split(' ')[0]) > 20:
            continue
    except:
        continue

    directory = os.path.join('acts',doc['status'].lower(), doc['title'][0].lower())
    filename = doc['title'].lower() + '.md'
    fullpath = os.path.join(directory, filename)
    fullpath = "".join(c for c in fullpath if c.isalnum() or c in keepcharacters).rstrip()
    converter.check_dir(directory)

    if WRITE and os.path.exists(fullpath):
        print 'already converted {}'.format(fullpath)
        continue

    if doc['type'] not in ['iconDOC', 'iconDOCX', 'iconRTF']:
        message = 'Unable to convert {} of type {}'.format(doc['title'], doc['type'])
        with open(fullpath, 'wb') as out:
            out.write(message)
        print message
        continue

    joined = converter.convert(input=os.path.join('comlaw',doc['uuid']), output=fullpath, type=doc['type'])
    joined = join_styles(joined)
    joined = apply_bold_italic(joined)
    joined = join_styles(joined)
    joined = convert_indentation(joined)
    result = apply_indentation(joined)
    if WRITE:
        print 'Converted {}'.format(fullpath)
        try:
            with open(fullpath, 'wb') as out:
                out.write(result)
        except:
            os.remove(fullpath)
            raise
    else:
        if result != '':
            print result




